"""
File Parser Service for extracting text from uploaded files.
Supports: .txt, .doc, .docx
"""
import logging
import os
import tempfile
from typing import Optional
from fastapi import UploadFile
from docx import Document

logger = logging.getLogger(__name__)


class FileParserService:
    """Service for parsing and extracting text from various file formats."""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.doc', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB in bytes
    
    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        """
        Extract text from uploaded file based on file type.
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Extracted text content as string
            
        Raises:
            ValueError: If file type is unsupported or file is empty
            Exception: If text extraction fails
        """
        # Get file extension
        filename = file.filename or ""
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Validate file extension
        if file_extension not in FileParserService.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(FileParserService.SUPPORTED_EXTENSIONS)}"
            )
        
        # Read file content
        content = await file.read()
        
        # Check if file is empty
        if not content or len(content) == 0:
            raise ValueError("The uploaded file is empty")
        
        # Check file size
        if len(content) > FileParserService.MAX_FILE_SIZE:
            raise ValueError(
                f"File size ({len(content)} bytes) exceeds maximum allowed size "
                f"({FileParserService.MAX_FILE_SIZE} bytes)"
            )
        
        # Extract text based on file type
        try:
            if file_extension == '.txt':
                return await FileParserService.extract_text_from_txt(content)
            elif file_extension in ['.doc', '.docx']:
                return await FileParserService.extract_text_from_docx(content, file_extension)
            else:
                raise ValueError(f"Unsupported file extension: {file_extension}")
                
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            raise Exception(f"Failed to extract text from file. The file may be corrupted: {str(e)}")
    
    @staticmethod
    async def extract_text_from_txt(content: bytes) -> str:
        """
        Extract text from .txt file.
        
        Args:
            content: File content as bytes
            
        Returns:
            Decoded text content
        """
        try:
            # Try UTF-8 first
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            try:
                text = content.decode('latin-1')
            except Exception as e:
                raise Exception(f"Failed to decode text file: {str(e)}")
        
        # Strip excessive whitespace but preserve line breaks
        text = text.strip()
        
        if not text:
            raise ValueError("The text file is empty or contains only whitespace")
        
        return text
    
    @staticmethod
    async def extract_text_from_docx(content: bytes, file_extension: str) -> str:
        """
        Extract text from .docx or .doc file.
        
        Args:
            content: File content as bytes
            file_extension: File extension (.doc or .docx)
            
        Returns:
            Extracted text content
        """
        # Create temporary file to save content
        temp_file_path = None
        
        try:
            # Create temporary file with appropriate extension
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(content)
                temp_file_path = tmp_file.name
            
            # Extract text using python-docx
            if file_extension == '.docx':
                doc = Document(temp_file_path)
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:  # Only add non-empty paragraphs
                        paragraphs.append(text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            paragraphs.append(" | ".join(row_text))
                
                # Join all paragraphs with double newline
                extracted_text = "\n\n".join(paragraphs)
                
            elif file_extension == '.doc':
                # For .doc files, we need textract or antiword
                # For MVP, we'll show an error message suggesting .docx
                raise ValueError(
                    "Legacy .doc format is not fully supported. "
                    "Please convert to .docx format or save as .txt"
                )
            else:
                raise ValueError(f"Unsupported Word document format: {file_extension}")
            
            if not extracted_text or not extracted_text.strip():
                raise ValueError("No text content found in the document")
            
            return extracted_text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
            
        finally:
            # Clean up temporary file
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {temp_file_path}: {e}")
